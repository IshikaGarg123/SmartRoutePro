import numpy as np
import pandas as pd
import networkx as nx
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.express as px
import streamlit as st
from plotly.subplots import make_subplots
import time
import heapq
from collections import defaultdict, deque
import random
from typing import Dict, List, Tuple, Optional, Set
import folium
from streamlit_folium import st_folium
import requests
import json
from datetime import datetime, timedelta
import base64
import io
import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pickle
from geopy.geocoders import Nominatim
from geopy.distance import geodesic
import warnings
warnings.filterwarnings('ignore')

# Set page configuration
st.set_page_config(
    page_title="SmartRoute Pro - Punjab Edition",
    page_icon="🚚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced Custom CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    .main-header {
        background: linear-gradient(135deg, #ff6b35 0%, #f7931e 100%);
        padding: 2.5rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 25px rgba(255, 107, 53, 0.3);
    }
    
    .main-header h1 {
        font-size: 3rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    
    .sidebar-section {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
        border: 1px solid #e2e8f0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .sidebar-header {
        font-size: 1.2rem;
        font-weight: 600;
        color: #ff6b35;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #fff 0%, #f8fafc 100%);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        margin: 0.5rem 0;
        text-align: center;
        transition: all 0.3s ease;
    }
    
    .algorithm-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
        border-left: 4px solid;
        margin: 1rem 0;
        transition: all 0.3s ease;
    }
    
    .dijkstra-card { border-color: #ff6b35; }
    .bellman-card { border-color: #10b981; }
    .astar-card { border-color: #3b82f6; }
    
    .weather-card {
        background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        text-align: center;
    }
    
    .traffic-indicator {
        display: inline-block;
        width: 12px;
        height: 12px;
        border-radius: 50%;
        margin-right: 8px;
    }
    
    .traffic-low { background-color: #10b981; }
    .traffic-medium { background-color: #f59e0b; }
    .traffic-high { background-color: #ef4444; }
    
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(2, 1fr);
        gap: 0.5rem;
        margin: 1rem 0;
    }
    
    .stats-item {
        background: white;
        padding: 0.75rem;
        border-radius: 6px;
        text-align: center;
        border: 1px solid #e2e8f0;
    }
    
    .punjab-badge {
        display: inline-block;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.875rem;
        margin: 0.25rem;
        background: linear-gradient(135deg, #ff6b35 0%, #f7931e 100%);
        color: white;
    }
    
    .success-box {
        background: linear-gradient(135deg, #10b981 0%, #34d399 100%);
        padding: 1.5rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        box-shadow: 0 5px 15px rgba(16, 185, 129, 0.3);
    }
    
    .custom-metric {
        text-align: center;
        padding: 1rem;
        background: white;
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        margin: 0.5rem 0;
    }
    
    .custom-metric .value {
        font-size: 1.5rem;
        font-weight: 700;
        color: #ff6b35;
    }
    
    .custom-metric .label {
        font-size: 0.9rem;
        color: #6b7280;
        margin-top: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

# Enhanced Punjab Cities with more details
PUNJAB_CITIES = {
    "Ludhiana": {"lat": 30.9010, "lon": 75.8573, "type": "major", "population": 1618000, "area": 310, "elevation": 244},
    "Amritsar": {"lat": 31.6340, "lon": 74.8723, "type": "major", "population": 1132761, "area": 286, "elevation": 234},
    "Jalandhar": {"lat": 31.3260, "lon": 75.5762, "type": "major", "population": 873725, "area": 200, "elevation": 228},
    "Patiala": {"lat": 30.3398, "lon": 76.3869, "type": "major", "population": 446246, "area": 95, "elevation": 250},
    "Chandigarh": {"lat": 30.7333, "lon": 76.7794, "type": "major", "population": 1025682, "area": 114, "elevation": 321},
    "Mohali": {"lat": 30.7046, "lon": 76.7179, "type": "major", "population": 146213, "area": 167, "elevation": 310},
    "Bathinda": {"lat": 30.2110, "lon": 74.9455, "type": "major", "population": 285813, "area": 97, "elevation": 211},
    "Firozpur": {"lat": 30.9320, "lon": 74.6150, "type": "medium", "population": 110313, "area": 55, "elevation": 204},
    "Hoshiarpur": {"lat": 31.5385, "lon": 75.9107, "type": "medium", "population": 168143, "area": 45, "elevation": 296},
    "Kapurthala": {"lat": 31.3800, "lon": 75.3800, "type": "medium", "population": 98916, "area": 35, "elevation": 225},
    "Moga": {"lat": 30.8081, "lon": 75.1708, "type": "medium", "population": 159897, "area": 40, "elevation": 217},
    "Pathankot": {"lat": 32.2746, "lon": 75.6522, "type": "medium", "population": 174683, "area": 87, "elevation": 332},
    "Rajpura": {"lat": 30.4779, "lon": 76.5946, "type": "medium", "population": 89263, "area": 25, "elevation": 259},
    "Sangrur": {"lat": 30.2458, "lon": 75.8421, "type": "medium", "population": 92347, "area": 45, "elevation": 237},
    "Tarn Taran": {"lat": 31.4515, "lon": 74.9289, "type": "small", "population": 65457, "area": 30, "elevation": 218},
    "Fatehgarh Sahib": {"lat": 30.6442, "lon": 76.3936, "type": "small", "population": 65064, "area": 20, "elevation": 249},
    "Mansa": {"lat": 29.9988, "lon": 75.3932, "type": "small", "population": 83437, "area": 35, "elevation": 207},
    "Gurdaspur": {"lat": 32.0409, "lon": 75.4046, "type": "medium", "population": 89498, "area": 40, "elevation": 265},
    "Fazilka": {"lat": 30.4034, "lon": 74.0244, "type": "small", "population": 76492, "area": 25, "elevation": 176},
    "Muktsar": {"lat": 30.4762, "lon": 74.5142, "type": "small", "population": 117085, "area": 30, "elevation": 196},
    "Faridkot": {"lat": 30.6728, "lon": 74.7348, "type": "small", "population": 91238, "area": 28, "elevation": 201},
    "Malout": {"lat": 30.2072, "lon": 74.4857, "type": "small", "population": 143264, "area": 22, "elevation": 193},
    "Khanna": {"lat": 30.7056, "lon": 76.2197, "type": "small", "population": 128137, "area": 18, "elevation": 245},
    "Phagwara": {"lat": 31.2212, "lon": 75.7839, "type": "small", "population": 100148, "area": 20, "elevation": 236},
    "Abohar": {"lat": 30.1204, "lon": 74.1995, "type": "small", "population": 145302, "area": 35, "elevation": 186},
}

# Vehicle types with specifications
VEHICLE_TYPES = {
    "Car": {"speed": 70, "fuel_efficiency": 15, "cost_per_km": 8, "capacity": 5, "icon": "🚗"},
    "Motorcycle": {"speed": 60, "fuel_efficiency": 45, "cost_per_km": 3, "capacity": 2, "icon": "🏍️"},
    "Bus": {"speed": 50, "fuel_efficiency": 4, "cost_per_km": 25, "capacity": 50, "icon": "🚌"},
    "Truck": {"speed": 45, "fuel_efficiency": 6, "cost_per_km": 35, "capacity": 10000, "icon": "🚛"},
    "Bicycle": {"speed": 15, "fuel_efficiency": 0, "cost_per_km": 0, "capacity": 1, "icon": "🚴"},
    "Auto-Rickshaw": {"speed": 40, "fuel_efficiency": 20, "cost_per_km": 12, "capacity": 3, "icon": "🛺"}
}
# Fuel types with specifications
FUEL_TYPES = {
    "Petrol": {
        "price_per_liter": 100,  # ₹ per liter
        "co2_per_liter": 2.31,   # kg CO2 per liter
        "compatible_vehicles": ["Car", "Motorcycle", "Auto-Rickshaw"],
        "icon": "⛽"
    },
    "Diesel": {
        "price_per_liter": 90,   # ₹ per liter
        "co2_per_liter": 2.68,   # kg CO2 per liter
        "compatible_vehicles": ["Car", "Bus", "Truck"],
        "icon": "🛢️"
    },
    "CNG": {
        "price_per_liter": 75,   # ₹ per kg (equivalent)
        "co2_per_liter": 1.85,   # kg CO2 per kg
        "compatible_vehicles": ["Car", "Bus", "Auto-Rickshaw"],
        "icon": "💨"
    },
    "Electric": {
        "price_per_liter": 8,    # ₹ per kWh (cost per km equivalent)
        "co2_per_liter": 0.5,    # kg CO2 per kWh (grid emissions)
        "compatible_vehicles": ["Car", "Bus", "Motorcycle", "Auto-Rickshaw"],
        "icon": "⚡"
    },
    "Bicycle": {
        "price_per_liter": 0,
        "co2_per_liter": 0,
        "compatible_vehicles": ["Bicycle"],
        "icon": "🚴"
    }
}
# Route preferences
ROUTE_PREFERENCES = {
    "Fastest": {"priority": "time", "description": "Minimize travel time", "icon": "⚡"},
    "Shortest": {"priority": "distance", "description": "Minimize total distance", "icon": "📏"},
    "Economical": {"priority": "cost", "description": "Minimize fuel/travel cost", "icon": "💰"},
    "Scenic": {"priority": "scenic", "description": "Pass through major cities", "icon": "🌄"},
    "Safe": {"priority": "safety", "description": "Use major highways only", "icon": "🛡️"}
}

class AdvancedPunjabPathfinder:
    """Advanced pathfinding with multiple optimization criteria"""
    
    def __init__(self):
        self.performance_data = []
        self.algorithm_stats = {
            'dijkstra': {'runs': 0, 'total_time': 0, 'avg_time': 0, 'success_rate': 0, 'successes': 0},
            'bellman_ford': {'runs': 0, 'total_time': 0, 'avg_time': 0, 'success_rate': 0, 'successes': 0},
            'a_star': {'runs': 0, 'total_time': 0, 'avg_time': 0, 'success_rate': 0, 'successes': 0},
            'floyd_warshall': {'runs': 0, 'total_time': 0, 'avg_time': 0, 'success_rate': 0, 'successes': 0}
        }
        self.route_history = []
    
    def calculate_haversine_distance(self, lat1: float, lon1: float, lat2: float, lon2: float) -> float:
        """Calculate the great circle distance between two points in kilometers"""
        R = 6371  # Earth's radius in kilometers
        lat1_rad = math.radians(lat1)
        lat2_rad = math.radians(lat2)
        delta_lat = math.radians(lat2 - lat1)
        delta_lon = math.radians(lon2 - lon1)
        
        a = math.sin(delta_lat/2)**2 + math.cos(lat1_rad) * math.cos(lat2_rad) * math.sin(delta_lon/2)**2
        c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
        distance = R * c
        return distance
    
    def calculate_route_cost(self, path: List[str], distances: Dict, vehicle_type: str, fuel_type: str = "Petrol") -> Tuple[float, float, float]:
        """Calculate total route cost based on vehicle type and fuel type"""
        if not path or len(path) < 2:
            return 0, 0, 0
    
        total_distance = 0
        for i in range(len(path) - 1):
            segment_dist = distances.get(path[i], {}).get(path[i + 1], 0)
            total_distance += segment_dist
    
        vehicle_specs = VEHICLE_TYPES.get(vehicle_type, VEHICLE_TYPES["Car"])
        fuel_specs = FUEL_TYPES.get(fuel_type, FUEL_TYPES["Petrol"])
    
        # Calculate fuel consumption
        if vehicle_specs["fuel_efficiency"] > 0:
            if fuel_type == "Electric":
            # For electric, calculate based on kWh per km
                fuel_consumption = total_distance * 0.15  # Average 0.15 kWh per km
            else:
                fuel_consumption = total_distance / vehicle_specs["fuel_efficiency"]
        else:
                fuel_consumption = 0
    
        # Calculate fuel cost
        fuel_cost = fuel_consumption * fuel_specs["price_per_liter"]
    
        # Calculate CO2 emissions
        co2_emissions = fuel_consumption * fuel_specs["co2_per_liter"]
    
        return fuel_cost, fuel_consumption, co2_emissions
    
    def calculate_travel_time(self, path: List[str], distances: Dict, vehicle_type: str, traffic_factor: float = 1.0) -> float:
        """Calculate travel time considering vehicle type and traffic"""
        if not path or len(path) < 2:
            return 0
        
        total_distance = 0
        for i in range(len(path) - 1):
            segment_dist = distances.get(path[i], {}).get(path[i + 1], 0)
            total_distance += segment_dist
        
        vehicle_specs = VEHICLE_TYPES.get(vehicle_type, VEHICLE_TYPES["Car"])
        avg_speed = vehicle_specs["speed"] * traffic_factor
        if avg_speed <= 0:
            return float('inf')
        travel_time_hours = total_distance / avg_speed
        return travel_time_hours * 60  # Return in minutes
    
    def dijkstra_multi_criteria(self, graph: Dict, start: str, end: str, 
                               vehicle_type: str = "Car", 
                               preference: str = "Fastest") -> Tuple[Dict, Dict, float, List]:
        """Enhanced Dijkstra with multi-criteria optimization"""
        start_time = time.perf_counter()
        
        def get_edge_weight(dist: float, from_city: str, to_city: str) -> float:
            vehicle_specs = VEHICLE_TYPES.get(vehicle_type, VEHICLE_TYPES["Car"])
            
            if preference == "Fastest":
                return dist / vehicle_specs["speed"]
            elif preference == "Shortest":
                return dist
            elif preference == "Economical":
                return dist * vehicle_specs["cost_per_km"]
            elif preference == "Scenic":
                city_bonus = 0
                if PUNJAB_CITIES[to_city]["type"] == "major":
                    city_bonus = -5
                return dist + city_bonus
            elif preference == "Safe":
                safety_factor = 1.0
                if PUNJAB_CITIES[from_city]["type"] == "major" and PUNJAB_CITIES[to_city]["type"] == "major":
                    safety_factor = 0.8
                return dist * safety_factor
            else:
                return dist
        
        distances = {city: float('inf') for city in graph}
        distances[start] = 0
        previous = {}
        pq = [(0, start)]
        visited = set()
        nodes_visited = []
        
        while pq:
            current_dist, current = heapq.heappop(pq)
            
            if current in visited:
                continue
                
            visited.add(current)
            nodes_visited.append(current)
            
            if current == end:
                break
                
            for neighbor, base_distance in graph[current].items():
                edge_weight = get_edge_weight(base_distance, current, neighbor)
                distance = current_dist + edge_weight
                
                if distance < distances[neighbor]:
                    distances[neighbor] = distance
                    previous[neighbor] = current
                    heapq.heappush(pq, (distance, neighbor))
        
        execution_time = (time.perf_counter() - start_time) * 1000
        self._update_stats('dijkstra', execution_time, distances[end] != float('inf'))
        return distances, previous, execution_time, nodes_visited
    
    def a_star_enhanced(self, graph: Dict, start: str, end: str, city_coords: Dict, 
                       vehicle_type: str = "Car", preference: str = "Fastest") -> Tuple[Dict, Dict, float, List]:
        """Enhanced A* with multi-criteria heuristic"""
        start_time = time.perf_counter()
        
        def heuristic(city1: str, city2: str) -> float:
            if city1 in city_coords and city2 in city_coords:
                h_distance = self.calculate_haversine_distance(
                    city_coords[city1]["lat"], city_coords[city1]["lon"],
                    city_coords[city2]["lat"], city_coords[city2]["lon"]
                )
                
                vehicle_specs = VEHICLE_TYPES.get(vehicle_type, VEHICLE_TYPES["Car"])
                
                if preference == "Fastest":
                    return h_distance / vehicle_specs["speed"]
                elif preference == "Economical":
                    return h_distance * vehicle_specs["cost_per_km"]
                else:
                    return h_distance
            return 0
        
        def get_edge_cost(dist: float, from_city: str, to_city: str) -> float:
            vehicle_specs = VEHICLE_TYPES.get(vehicle_type, VEHICLE_TYPES["Car"])
            
            if preference == "Fastest":
                return dist / vehicle_specs["speed"]
            elif preference == "Economical":
                return dist * vehicle_specs["cost_per_km"]
            elif preference == "Scenic":
                city_bonus = 0
                if PUNJAB_CITIES[to_city]["type"] == "major":
                    city_bonus = -3
                return dist + city_bonus
            else:
                return dist
        
        open_set = [(0, start)]
        g_score = {city: float('inf') for city in graph}
        g_score[start] = 0
        f_score = {city: float('inf') for city in graph}
        f_score[start] = heuristic(start, end)
        previous = {}
        nodes_visited = []
        closed_set = set()
        
        while open_set:
            current_f, current = heapq.heappop(open_set)
            
            if current in closed_set:
                continue
                
            closed_set.add(current)
            nodes_visited.append(current)
            
            if current == end:
                break
                
            for neighbor, base_distance in graph[current].items():
                if neighbor in closed_set:
                    continue
                
                tentative_g = g_score[current] + get_edge_cost(base_distance, current, neighbor)
                
                if tentative_g < g_score[neighbor]:
                    previous[neighbor] = current
                    g_score[neighbor] = tentative_g
                    f_score[neighbor] = g_score[neighbor] + heuristic(neighbor, end)
                    heapq.heappush(open_set, (f_score[neighbor], neighbor))
        
        execution_time = (time.perf_counter() - start_time) * 1000
        self._update_stats('a_star', execution_time, g_score[end] != float('inf'))
        return g_score, previous, execution_time, nodes_visited
    
    def floyd_warshall_all_pairs(self, graph: Dict) -> Tuple[Dict, Dict, float]:
        """Floyd-Warshall for all-pairs shortest paths"""
        start_time = time.perf_counter()
        
        cities = list(graph.keys())
        dist = {}
        next_node = {}
        
        # Initialize
        for i in cities:
            dist[i] = {}
            next_node[i] = {}
            for j in cities:
                if i == j:
                    dist[i][j] = 0
                    next_node[i][j] = j
                elif j in graph[i]:
                    dist[i][j] = graph[i][j]
                    next_node[i][j] = j
                else:
                    dist[i][j] = float('inf')
                    next_node[i][j] = None
        
        # Main algorithm
        for k in cities:
            for i in cities:
                for j in cities:
                    if dist[i][k] + dist[k][j] < dist[i][j]:
                        dist[i][j] = dist[i][k] + dist[k][j]
                        next_node[i][j] = next_node[i][k]
        
        execution_time = (time.perf_counter() - start_time) * 1000
        self._update_stats('floyd_warshall', execution_time, True)
        return dist, next_node, execution_time
    
    def bellman_ford_enhanced(self, graph: Dict, start: str) -> Tuple[Dict, Dict, float, bool]:
        """Enhanced Bellman-Ford with early termination"""
        start_time = time.perf_counter()
        
        cities = set(graph.keys())
        for city in graph:
            cities.update(graph[city].keys())
        
        distances = {city: float('inf') for city in cities}
        distances[start] = 0
        previous = {}
        
        # Relax edges V-1 times
        for iteration in range(len(cities) - 1):
            updated = False
            for city in graph:
                if distances[city] != float('inf'):
                    for neighbor, weight in graph[city].items():
                        if distances[city] + weight < distances[neighbor]:
                            distances[neighbor] = distances[city] + weight
                            previous[neighbor] = city
                            updated = True
            
            if not updated:  # Early termination
                break
        
        # Check for negative cycles
        has_negative_cycle = False
        for city in graph:
            if distances[city] != float('inf'):
                for neighbor, weight in graph[city].items():
                    if distances[city] + weight < distances[neighbor]:
                        has_negative_cycle = True
                        break
        
        execution_time = (time.perf_counter() - start_time) * 1000
        self._update_stats('bellman_ford', execution_time, not has_negative_cycle)
        return distances, previous, execution_time, has_negative_cycle
    
    def _update_stats(self, algorithm: str, execution_time: float, success: bool):
        """Update algorithm statistics"""
        stats = self.algorithm_stats[algorithm]
        stats['runs'] += 1
        stats['total_time'] += execution_time
        stats['avg_time'] = stats['total_time'] / stats['runs']
        if success:
            stats['successes'] += 1
        stats['success_rate'] = stats['successes'] / stats['runs']
    
    def get_path(self, previous: Dict, start: str, end: str) -> List[str]:
        """Reconstruct path from previous nodes dictionary"""
        path = []
        current = end
        
        while current is not None:
            path.append(current)
            current = previous.get(current)
        
        path.reverse()
        return path if path and path[0] == start else []
    
    def save_route_history(self, source: str, destination: str, path: List[str], 
                          distance: float, vehicle_type: str, preference: str):
        """Save route to history"""
        route_data = {
            'timestamp': datetime.now(),
            'source': source,
            'destination': destination,
            'path': path,
            'distance': distance,
            'vehicle_type': vehicle_type,
            'preference': preference
        }
        self.route_history.append(route_data)

class PunjabNetworkGenerator:
    """Enhanced Punjab network generator with realistic connections"""
    
    @staticmethod
    def generate_enhanced_punjab_network() -> Dict:
        """Generate enhanced Punjab network with highway classifications"""
        graph = defaultdict(dict)
        
        # Major highway connections
        major_connections = {
            "Ludhiana": {"Jalandhar": 85, "Patiala": 65, "Chandigarh": 98, "Sangrur": 52, "Khanna": 45, "Moga": 45},
            "Amritsar": {"Jalandhar": 78, "Tarn Taran": 25, "Pathankot": 117, "Gurdaspur": 50, "Kapurthala": 95},
            "Jalandhar": {"Ludhiana": 85, "Amritsar": 78, "Kapurthala": 20, "Hoshiarpur": 45, "Phagwara": 25, "Patiala": 120},
            "Patiala": {"Ludhiana": 65, "Chandigarh": 65, "Sangrur": 45, "Rajpura": 25, "Fatehgarh Sahib": 12, "Jalandhar": 120},
            "Chandigarh": {"Patiala": 65, "Mohali": 8, "Ludhiana": 98, "Rajpura": 45, "Khanna": 78},
            "Mohali": {"Chandigarh": 8, "Rajpura": 40, "Khanna": 55, "Patiala": 70},
            "Bathinda": {"Mansa": 35, "Faridkot": 48, "Sangrur": 105, "Malout": 52, "Muktsar": 78, "Abohar": 85},
            "Firozpur": {"Faridkot": 32, "Muktsar": 45, "Fazilka": 65, "Ludhiana": 125},
            "Hoshiarpur": {"Jalandhar": 45, "Kapurthala": 55, "Phagwara": 65, "Pathankot": 95},
            "Kapurthala": {"Jalandhar": 20, "Hoshiarpur": 55, "Phagwara": 35, "Amritsar": 95},
            "Moga": {"Ludhiana": 45, "Faridkot": 38, "Bathinda": 78, "Firozpur": 55},
            "Pathankot": {"Amritsar": 117, "Gurdaspur": 75, "Hoshiarpur": 95},
            "Rajpura": {"Patiala": 25, "Chandigarh": 45, "Mohali": 40, "Khanna": 35},
            "Sangrur": {"Ludhiana": 52, "Patiala": 45, "Bathinda": 105, "Mansa": 75, "Malout": 88},
            "Tarn Taran": {"Amritsar": 25, "Gurdaspur": 65},
            "Fatehgarh Sahib": {"Patiala": 12, "Chandigarh": 55, "Khanna": 42},
            "Mansa": {"Bathinda": 35, "Sangrur": 75, "Malout": 45},
            "Gurdaspur": {"Amritsar": 50, "Pathankot": 75, "Tarn Taran": 65},
            "Fazilka": {"Firozpur": 65, "Abohar": 25, "Muktsar": 78},
            "Muktsar": {"Firozpur": 45, "Faridkot": 55, "Bathinda": 78, "Fazilka": 78},
            "Faridkot": {"Firozpur": 32, "Muktsar": 55, "Bathinda": 48, "Moga": 38},
            "Malout": {"Bathinda": 52, "Abohar": 45, "Mansa": 45, "Sangrur": 88},
            "Khanna": {"Ludhiana": 45, "Mohali": 55, "Patiala": 38, "Chandigarh": 78, "Rajpura": 35, "Fatehgarh Sahib": 42},
            "Phagwara": {"Jalandhar": 25, "Kapurthala": 35, "Hoshiarpur": 65},
            "Abohar": {"Fazilka": 25, "Malout": 45, "Bathinda": 85}
        }
        
        # Create bidirectional graph with highway classifications
        for city, neighbors in major_connections.items():
            for neighbor, distance in neighbors.items():
                # Add some variation for realistic road conditions
                road_condition_factor = random.uniform(0.95, 1.05)
                adjusted_distance = distance * road_condition_factor
                
                graph[city][neighbor] = round(adjusted_distance, 1)
                graph[neighbor][city] = round(adjusted_distance, 1)  # Bidirectional
        
        return dict(graph)
    
    @staticmethod
    def get_traffic_conditions() -> Dict:
        """Simulate current traffic conditions"""
        traffic_data = {}
        for city in PUNJAB_CITIES.keys():
            # Simulate traffic based on city type and random factors
            base_traffic = {
                "major": random.uniform(0.8, 1.3),
                "medium": random.uniform(0.9, 1.2),
                "small": random.uniform(0.95, 1.1)
            }
            
            city_type = PUNJAB_CITIES[city]["type"]
            traffic_factor = base_traffic[city_type]
            
            if traffic_factor > 1.2:
                traffic_level = "High"
                color = "traffic-high"
            elif traffic_factor > 1.1:
                traffic_level = "Medium"
                color = "traffic-medium"
            else:
                traffic_level = "Low"
                color = "traffic-low"
            
            traffic_data[city] = {
                "factor": traffic_factor,
                "level": traffic_level,
                "color": color
            }
        
        return traffic_data
    
    @staticmethod
    def get_weather_conditions() -> Dict:
        """Simulate weather conditions"""
        weather_conditions = ["Clear", "Partly Cloudy", "Cloudy", "Light Rain", "Heavy Rain", "Fog"]
        weather = random.choice(weather_conditions)
        
        weather_factors = {
            "Clear": 1.0,
            "Partly Cloudy": 1.0,
            "Cloudy": 0.95,
            "Light Rain": 0.8,
            "Heavy Rain": 0.6,
            "Fog": 0.7
        }
        
        return {
            "condition": weather,
            "factor": weather_factors[weather],
            "temperature": random.randint(15, 35),
            "humidity": random.randint(40, 90)
        }

def create_enhanced_punjab_map(source: str = None, destination: str = None, 
                              path: List[str] = None, algorithm_name: str = None,
                              traffic_data: Dict = None) -> folium.Map:
    """Create enhanced interactive Punjab map with traffic and weather info"""
    punjab_center = [30.7333, 75.5762]
    
    m = folium.Map(
        location=punjab_center,
        zoom_start=8,
        tiles='CartoDB positron',
        attr='Punjab Route Optimizer Pro'
    )
    
    # Add all Punjab cities with enhanced information
    for city, data in PUNJAB_CITIES.items():
        lat, lon = data["lat"], data["lon"]
        city_type = data["type"]
        population = data["population"]
        
        # Different colors and sizes based on city importance
        if city_type == "major":
            color = "#ff6b35"
            radius = 15
            fillColor = "#ff6b35"
        elif city_type == "medium":
            color = "#3b82f6"
            radius = 10
            fillColor = "#3b82f6"
        else:
            color = "#10b981"
            radius = 7
            fillColor = "#10b981"
        
        # Highlight source and destination
        if city == source:
            color = "#059669"
            fillColor = "#059669"
            radius = 18
        elif city == destination:
            color = "#dc2626"
            fillColor = "#dc2626"
            radius = 18
        
        # Add traffic indicator if available
        traffic_info = ""
        if traffic_data and city in traffic_data:
            traffic_level = traffic_data[city]["level"]
            traffic_info = f"<p><strong>Traffic:</strong> {traffic_level}</p>"
        
        # Create enhanced popup content
        popup_content = f"""
        <div style="font-family: Arial; width: 250px;">
            <h4 style="margin: 0; color: #ff6b35;">{city}</h4>
            <p style="margin: 5px 0;"><strong>Type:</strong> {city_type.title()} City</p>
            <p style="margin: 5px 0;"><strong>Population:</strong> {population:,}</p>
            <p style="margin: 5px 0;"><strong>Area:</strong> {data['area']} km²</p>
            <p style="margin: 5px 0;"><strong>Elevation:</strong> {data['elevation']} m</p>
            {traffic_info}
            <p style="margin: 5px 0;"><strong>Coordinates:</strong><br>{lat:.4f}°N, {lon:.4f}°E</p>
        </div>
        """
        
        folium.CircleMarker(
            location=[lat, lon],
            radius=radius,
            popup=folium.Popup(popup_content, max_width=280),
            color="white",
            weight=3,
            fillColor=fillColor,
            fillOpacity=0.8,
            tooltip=f"{city} (Pop: {population:,})"
        ).add_to(m)
    
    # Add route path with enhanced styling
    if path and len(path) > 1:
        route_coords = []
        for city in path:
            if city in PUNJAB_CITIES:
                route_coords.append([PUNJAB_CITIES[city]["lat"], PUNJAB_CITIES[city]["lon"]])
        
        if route_coords:
            # Add main route line
            folium.PolyLine(
                locations=route_coords,
                weight=6,
                color="#dc2626",
                opacity=0.9,
                popup=f"{algorithm_name} Route: {' → '.join(path)}" if algorithm_name else f"Route: {' → '.join(path)}"
            ).add_to(m)
    
    # Add enhanced legend
    legend_html = '''
    <div style="position: fixed; 
                top: 10px; right: 10px; width: 250px; height: auto; 
                background-color: white; border:2px solid #ff6b35; z-index:9999; 
                font-size:12px; padding: 15px; border-radius: 10px;
                box-shadow: 0 4px 12px rgba(0,0,0,0.15);">
    <h4 style="margin: 0 0 10px 0; color: #ff6b35;">Punjab Route Map</h4>
    <p style="margin: 3px 0;"><span style="color:#ff6b35;font-size:18px;">●</span> Major Cities</p>
    <p style="margin: 3px 0;"><span style="color:#3b82f6;font-size:14px;">●</span> Medium Cities</p>
    <p style="margin: 3px 0;"><span style="color:#10b981;font-size:12px;">●</span> Small Cities</p>
    <hr style="margin: 8px 0;">
    <p style="margin: 3px 0;"><span style="color:#059669;font-size:20px;">●</span> Source</p>
    <p style="margin: 3px 0;"><span style="color:#dc2626;font-size:20px;">●</span> Destination</p>
    <p style="margin: 3px 0;"><span style="color:#dc2626;">━━━</span> Optimal Route</p>
    </div>
    '''
    m.get_root().html.add_child(folium.Element(legend_html))
    
    return m

def create_comprehensive_analysis_dashboard(results: Dict, vehicle_type: str, preference: str) -> go.Figure:
    """Create comprehensive analysis dashboard"""
    algorithms = list(results.keys())
    times = [results[algo]['time'] for algo in algorithms]
    distances = [results[algo].get('distance', float('inf')) for algo in algorithms]
    costs = [results[algo].get('cost', 0) for algo in algorithms]
    
    # Filter out infinite values
    finite_distances = [d if d != float('inf') else 0 for d in distances]
    
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=(
            f'Execution Time (ms) - {vehicle_type}',
            f'Route Distance (km) - {preference} Priority',
            'Travel Cost Comparison (₹)',
            'Algorithm Performance'
        )
    )
    
    colors = ['#ff6b35', '#10b981', '#3b82f6', '#f59e0b']
    
    # Execution time
    fig.add_trace(
        go.Bar(
            x=algorithms, 
            y=times, 
            name='Time', 
            marker_color=colors[:len(algorithms)],
            text=[f'{t:.3f}ms' for t in times],
            textposition='outside'
        ),
        row=1, col=1
    )
    
    # Distance comparison
    fig.add_trace(
        go.Bar(
            x=algorithms, 
            y=finite_distances, 
            name='Distance', 
            marker_color=colors[:len(algorithms)],
            text=[f'{d:.1f} km' if d != 0 else 'No Path' for d in finite_distances],
            textposition='outside',
            showlegend=False
        ),
        row=1, col=2
    )
    
    # Cost comparison
    fig.add_trace(
        go.Bar(
            x=algorithms,
            y=costs,
            name='Cost',
            marker_color=colors[:len(algorithms)],
            text=[f'₹{c:.0f}' for c in costs],
            textposition='outside',
            showlegend=False
        ),
        row=2, col=1
    )
    
    # Performance metrics
    if len(algorithms) > 0:
        performance_data = []
        for algo in algorithms:
            perf_score = 100 - min(results[algo]['time'], 100)  # Simple performance score
            performance_data.append(perf_score)
        
        fig.add_trace(
            go.Scatter(
                x=algorithms,
                y=performance_data,
                mode='markers+lines',
                name='Performance Score',
                marker=dict(size=10, color=colors[:len(algorithms)]),
                showlegend=False
            ),
            row=2, col=2
        )
    
    fig.update_layout(
        height=800,
        title_text=f"Comprehensive Route Analysis - {vehicle_type} ({preference} Priority)",
        title_x=0.5
    )
    
    return fig

def main():
    # Enhanced Header
    st.markdown("""
    <div class="main-header">
        <h1>SmartRoute Pro - Punjab Edition</h1>
        <h3>Advanced Multi-Criteria Route Optimization Platform</h3>
        <p>Find optimal routes with real-time traffic, weather, and vehicle-specific optimization</p>
        <div style="margin-top: 1rem;">
            <span class="punjab-badge">25 Cities</span>
            <span class="punjab-badge">6 Vehicles</span>
            <span class="punjab-badge">4 Algorithms</span>
            <span class="punjab-badge">5 Preferences</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize enhanced pathfinder
    pathfinder = AdvancedPunjabPathfinder()
    
    # Generate network and conditions
    if 'punjab_network' not in st.session_state:
        st.session_state.punjab_network = PunjabNetworkGenerator.generate_enhanced_punjab_network()
        st.session_state.traffic_data = PunjabNetworkGenerator.get_traffic_conditions()
        st.session_state.weather_data = PunjabNetworkGenerator.get_weather_conditions()
    
    # Enhanced Sidebar
    with st.sidebar:
        st.markdown("# 🚚 SmartRoute Pro")
        st.markdown("---")
        
        # City Selection Section
        st.markdown("## 📍 Route Configuration")
        
        st.markdown("### Select Cities")
        cities_list = sorted(PUNJAB_CITIES.keys())
        
        source_city = st.selectbox(
            "🟢 Source City",
            cities_list,
            index=cities_list.index("Ludhiana"),
            help="Starting point of your journey"
        )
        
        destination_city = st.selectbox(
            "🔴 Destination City", 
            cities_list,
            index=cities_list.index("Amritsar"),
            help="Your final destination"
        )
        
        # Quick city info
        if source_city != destination_city:
            source_info = PUNJAB_CITIES[source_city]
            dest_info = PUNJAB_CITIES[destination_city]
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Source Pop.", f"{source_info['population']:,}")
            
            with col_b:
                st.metric("Dest. Pop.", f"{dest_info['population']:,}")
        
        st.markdown("---")
        
        # Vehicle Selection Section
        st.markdown("## 🚗 Vehicle Selection")
        
        vehicle_type = st.selectbox(
            "Choose Vehicle Type",
            list(VEHICLE_TYPES.keys()),
            index=0,
            format_func=lambda x: f"{VEHICLE_TYPES[x]['icon']} {x}",
            help="Vehicle affects speed, cost, and route preferences"
        )
        st.markdown("---")
        
        # Fuel Type Selection Section
        st.markdown("## ⛽ Fuel Type")
        
        # Filter fuel types compatible with selected vehicle
        compatible_fuels = [fuel for fuel, specs in FUEL_TYPES.items() 
                          if vehicle_type in specs["compatible_vehicles"]]
        
        if compatible_fuels:
            fuel_type = st.selectbox(
                "Choose Fuel Type",
                compatible_fuels,
                index=0,
                format_func=lambda x: f"{FUEL_TYPES[x]['icon']} {x}",
                help="Fuel type affects cost and environmental impact"
            )
            
            # Fuel specifications display
            fuel_specs = FUEL_TYPES[fuel_type]
            
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                st.metric("Price/Unit", f"₹{fuel_specs['price_per_liter']}")
            
            with col_f2:
                st.metric("CO2/Unit", f"{fuel_specs['co2_per_liter']:.2f} kg")
        else:
            fuel_type = "Petrol"
            st.warning("No compatible fuel types for this vehicle")
            
        # Vehicle specifications display
        vehicle_specs = VEHICLE_TYPES[vehicle_type]
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Speed", f"{vehicle_specs['speed']} km/h")
            st.metric("Cost/km", f"₹{vehicle_specs['cost_per_km']}")
        
        with col2:
            st.metric("Fuel Eff.", f"{vehicle_specs['fuel_efficiency']} km/l")
            st.metric("Capacity", f"{vehicle_specs['capacity']} {'kg' if vehicle_type == 'Truck' else 'persons'}")
        
        st.markdown("---")
        
        # Route Preferences Section
        st.markdown("## 🎯 Route Preferences")
        
        preference = st.selectbox(
            "Optimization Priority",
            list(ROUTE_PREFERENCES.keys()),
            index=0,
            format_func=lambda x: f"{ROUTE_PREFERENCES[x]['icon']} {x}",
            help="What to optimize for in route calculation"
        )
        
        st.info(ROUTE_PREFERENCES[preference]['description'])
        
        st.markdown("---")
        
        # Algorithm Selection Section
        st.markdown("## ⚙️ Algorithm Selection")
        
        selected_algorithms = st.multiselect(
            "Choose Algorithms",
            ["Dijkstra", "A* (A-Star)", "Bellman-Ford", "Floyd-Warshall"],
            default=["Dijkstra", "A* (A-Star)"],
            help="Select algorithms for comparison"
        )
        
        st.markdown("---")
        
        # Real-time Conditions Section
        st.markdown("## 🌤️ Current Conditions")
        
        # Weather display
        weather = st.session_state.weather_data
        st.markdown(f"""
        **Weather:** {weather['condition']}  
        **Temperature:** {weather['temperature']}°C  
        **Humidity:** {weather['humidity']}%  
        **Speed Impact:** {weather['factor']*100:.0f}%
        """)
        
        # Traffic conditions for selected cities
        if source_city != destination_city:
            traffic = st.session_state.traffic_data
            
            st.markdown("**Traffic Conditions:**")
            
            source_traffic = traffic.get(source_city, {"level": "Unknown"})
            dest_traffic = traffic.get(destination_city, {"level": "Unknown"})
            
            col_t1, col_t2 = st.columns(2)
            with col_t1:
                st.markdown(f"**{source_city}:**  \n{source_traffic['level']}")
            
            with col_t2:
                st.markdown(f"**{destination_city}:**  \n{dest_traffic['level']}")
        
        st.markdown("---")
        
        # Action Buttons Section
        st.markdown("## 🚀 Actions")
        
        calculate_route = st.button("🔍 Calculate Route", type="primary", use_container_width=True)
        
        refresh_conditions = st.button("🔄 Refresh Conditions", use_container_width=True)
        
        st.markdown("---")
        
        # Quick Stats
        st.markdown("## 📊 Network Info")
        st.markdown(f"""
        **Cities:** {len(PUNJAB_CITIES)}  
        **Algorithms:** {len(selected_algorithms)} selected  
        **Vehicle:** {vehicle_type}  
        **Priority:** {preference}
        """)
    
    # Refresh conditions
    if refresh_conditions:
        st.session_state.traffic_data = PunjabNetworkGenerator.get_traffic_conditions()
        st.session_state.weather_data = PunjabNetworkGenerator.get_weather_conditions()
        st.rerun()
    
    # Main content area
    if source_city == destination_city:
        st.warning("Please select different source and destination cities.")
        punjab_map = create_enhanced_punjab_map(traffic_data=st.session_state.traffic_data)
        st_folium(punjab_map, width=700, height=500)
    
    elif calculate_route and selected_algorithms:
        with st.spinner(f"Calculating optimal {preference.lower()} route from {source_city} to {destination_city}..."):
            results = {}
            network = st.session_state.punjab_network
            traffic_data = st.session_state.traffic_data
            weather_data = st.session_state.weather_data
            
            # Apply traffic and weather factors
            traffic_factor = 1.0
            if source_city in traffic_data:
                traffic_factor *= traffic_data[source_city]["factor"]
            
            weather_factor = weather_data["factor"]
            combined_factor = traffic_factor * weather_factor
            
            # Progress tracking
            progress_bar = st.progress(0)
            total_algorithms = len(selected_algorithms)
            
            for i, algorithm in enumerate(selected_algorithms):
                try:
                    if algorithm == "Dijkstra":
                        distances, previous, time_taken, nodes_visited = pathfinder.dijkstra_multi_criteria(
                            network, source_city, destination_city, vehicle_type, preference
                        )
                        path = pathfinder.get_path(previous, source_city, destination_city)
                        
                        route_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
                            path, network, vehicle_type, fuel_type
                        )
                        travel_time = pathfinder.calculate_travel_time(path, network, vehicle_type, combined_factor)
                        
                        results["Dijkstra"] = {
                            'time': time_taken,
                            'distance': distances.get(destination_city, float('inf')),
                            'path': path,
                            'nodes_visited': nodes_visited,
                            'cost': route_cost,
                            'travel_time': travel_time,
                            'fuel_consumption': fuel_consumption,
                            'co2_emissions': co2_emissions
                        }
                        
                    
                    elif algorithm == "A* (A-Star)":
                        distances, previous, time_taken, nodes_visited = pathfinder.a_star_enhanced(
                            network, source_city, destination_city, PUNJAB_CITIES, vehicle_type, preference
                        )
                        path = pathfinder.get_path(previous, source_city, destination_city)
                        
                        route_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
                            path, network, vehicle_type, fuel_type
                        )
                        travel_time = pathfinder.calculate_travel_time(path, network, vehicle_type, combined_factor)
                        
                        results["A* (A-Star)"] = {
                            'time': time_taken,
                            'distance': distances.get(destination_city, float('inf')),
                            'path': path,
                            'nodes_visited': nodes_visited,
                            'cost': route_cost,
                            'travel_time': travel_time,
                            'fuel_consumption': fuel_consumption,
                            'co2_emissions': co2_emissions
                        }
                    
                    elif algorithm == "Bellman-Ford":
                        distances, previous, time_taken, has_neg_cycle = pathfinder.bellman_ford_enhanced(
                            network, source_city
                        )
                        path = pathfinder.get_path(previous, source_city, destination_city)
                        
                        route_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
                            path, network, vehicle_type, fuel_type
                        )
                        travel_time = pathfinder.calculate_travel_time(path, network, vehicle_type, combined_factor)
                        
                        results["Bellman-Ford"] = {
                            'time': time_taken,
                            'distance': distances.get(destination_city, float('inf')),
                            'path': path,
                            'nodes_visited': nodes_visited,
                            'cost': route_cost,
                            'travel_time': travel_time,
                            'fuel_consumption': fuel_consumption,
                            'co2_emissions': co2_emissions
                        }
                    
                    elif algorithm == "Floyd-Warshall":
                        dist_matrix, next_matrix, time_taken = pathfinder.floyd_warshall_all_pairs(network)
                        distance = dist_matrix.get(source_city, {}).get(destination_city, float('inf'))
                        
                        # Reconstruct path for Floyd-Warshall
                        path = []
                        if distance != float('inf'):
                            current = source_city
                            path = [current]
                            while current != destination_city:
                                current = next_matrix[current][destination_city]
                                if current is None:
                                    path = []
                                    break
                                path.append(current)
                        
                        route_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
                            path, network, vehicle_type, fuel_type
                        )
                        travel_time = pathfinder.calculate_travel_time(path, network, vehicle_type, combined_factor)
                        
                        results["Floydd-Warshall"] = {
                            'time': time_taken,
                            'distance': distances.get(destination_city, float('inf')),
                            'path': path,
                            'nodes_visited': nodes_visited,
                            'cost': route_cost,
                            'travel_time': travel_time,
                            'fuel_consumption': fuel_consumption,
                            'co2_emissions': co2_emissions
                        }
                    
                    progress_bar.progress((i + 1) / total_algorithms)
                    
                except Exception as e:
                    st.error(f"Error with {algorithm}: {str(e)}")
            
            progress_bar.empty()
            
            # Save results
            st.session_state.route_results = results
            st.session_state.current_source = source_city
            st.session_state.current_destination = destination_city
            st.session_state.current_vehicle = vehicle_type
            st.session_state.current_preference = preference
    
    # Display results if available
    if ('route_results' in st.session_state and 
        st.session_state.get('current_source') == source_city and 
        st.session_state.get('current_destination') == destination_city):
        
        results = st.session_state.route_results
        valid_results = {k: v for k, v in results.items() if v.get('distance', float('inf')) != float('inf')}
        network = st.session_state.punjab_network  # Make network available for results display
        
        if valid_results:
            # Find best routes
            best_distance_algo = min(valid_results.keys(), key=lambda x: valid_results[x]['distance'])
            best_path = valid_results[best_distance_algo]['path']
            best_distance = valid_results[best_distance_algo]['distance']
            
            # Create tabs
            tab1, tab2, tab3 = st.tabs([
                "Interactive Route Map", 
                "Performance Dashboard", 
                "Detailed Analysis"
            ])
            
            with tab1:
                st.markdown(f"### Optimal Route: {source_city} → {destination_city}")
                
                best_result = valid_results[best_distance_algo]
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Best Algorithm", best_distance_algo)
                with col2:
                    st.metric("Distance", f"{best_distance:.1f} km")
                with col3:
                    st.metric("Travel Time", f"{best_result['travel_time']:.0f} min")
                with col4:
                    st.metric("Estimated Cost", f"₹{best_result['cost']:.0f}")
                
                # Enhanced interactive map
                route_map = create_enhanced_punjab_map(
                    source_city, destination_city, best_path, 
                    best_distance_algo, st.session_state.traffic_data
                )
                st_folium(route_map, width=700, height=500)
                
                # Route table
                if best_path:
                    st.markdown("### Step-by-Step Route")
                    route_data = []
                    cumulative_distance = 0
                    
                    for i, city in enumerate(best_path):
                        city_info = PUNJAB_CITIES[city]
                        
                        if i == 0:
                            route_data.append({
                                "Step": i + 1,
                                "City": city,
                                "Type": city_info["type"].title(),
                                "Population": f"{city_info['population']:,}",
                                "Distance": "START",
                                "Cumulative": 0
                            })
                        else:
                            segment_dist = network.get(best_path[i-1], {}).get(city, 0)
                            cumulative_distance += segment_dist
                            
                            route_data.append({
                                "Step": i + 1,
                                "City": city,
                                "Type": city_info["type"].title(),
                                "Population": f"{city_info['population']:,}",
                                "Distance": f"{segment_dist:.1f} km",
                                "Cumulative": f"{cumulative_distance:.1f} km"
                            })
                    
                    df = pd.DataFrame(route_data)
                    st.dataframe(df, use_container_width=True)
            
            with tab2:
                st.markdown("### Performance Analysis")
                
                # Algorithm comparison
                cols = st.columns(min(len(results), 4))
                
                for i, (algo_name, data) in enumerate(results.items()):
                    with cols[i % 4]:
                        card_class = "dijkstra-card" if algo_name == "Dijkstra" else "astar-card" if algo_name == "A* (A-Star)" else "bellman-card"
                        
                        st.markdown(f'<div class="algorithm-card {card_class}">', unsafe_allow_html=True)
                        st.markdown(f"**{algo_name}**")
                        
                        if data.get('distance', float('inf')) == float('inf'):
                            st.error("No path found")
                        else:
                            col_a, col_b = st.columns(2)
                            with col_a:
                                st.metric("Time", f"{data['time']:.3f}ms")
                            with col_b:
                                st.metric("Distance", f"{data['distance']:.1f}km")
                            
                            st.metric("Cost", f"₹{data.get('cost', 0):.0f}")
                            st.metric("Travel Time", f"{data.get('travel_time', 0):.0f}min")
                        
                        st.markdown('</div>', unsafe_allow_html=True)
                
                # Dashboard
                if len(valid_results) > 1:
                    dashboard_fig = create_comprehensive_analysis_dashboard(
                        valid_results, vehicle_type, preference
                    )
                    st.plotly_chart(dashboard_fig, use_container_width=True)
            
            with tab3:
                st.markdown("### Detailed Analysis")
                
                # Route efficiency metrics
                col1, col2 = st.columns(2)
                
                with col1:
                    # Distance vs Direct distance
                    direct_distance = pathfinder.calculate_haversine_distance(
                        PUNJAB_CITIES[source_city]["lat"], PUNJAB_CITIES[source_city]["lon"],
                        PUNJAB_CITIES[destination_city]["lat"], PUNJAB_CITIES[destination_city]["lon"]
                    )
                    
                    efficiency = (direct_distance / best_distance) * 100 if best_distance > 0 else 0
                    detour_distance = best_distance - direct_distance
                    
                    st.markdown(f"""
                    <div class="success-box">
                        <h4>Route Efficiency Analysis</h4>
                        <p><strong>Direct Distance:</strong> {direct_distance:.1f} km</p>
                        <p><strong>Route Distance:</strong> {best_distance:.1f} km</p>
                        <p><strong>Route Efficiency:</strong> {efficiency:.1f}%</p>
                        <p><strong>Extra Distance:</strong> {detour_distance:.1f} km</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    # City types in route
                    if best_path:
                        city_types_count = {"major": 0, "medium": 0, "small": 0}
                        for city in best_path:
                            city_type = PUNJAB_CITIES[city]["type"]
                            city_types_count[city_type] += 1
                        
                        # City types pie chart
                        fig_pie = px.pie(
                            values=list(city_types_count.values()),
                            names=[f"{k.title()} Cities" for k in city_types_count.keys()],
                            title="City Types in Route"
                        )
                        st.plotly_chart(fig_pie, use_container_width=True)
                
                # Environmental impact
                if best_path and vehicle_type in VEHICLE_TYPES:
                    vehicle_specs = VEHICLE_TYPES[vehicle_type]
                    fuel_consumption = best_distance / vehicle_specs['fuel_efficiency'] if vehicle_specs['fuel_efficiency'] > 0 else 0
                    co2_emissions = fuel_consumption * 2.31  # kg CO2 per liter of petrol
                    
                    st.markdown("### Environmental Impact")
                    col_env1, col_env2, col_env3, col_env4 = st.columns(4)
                    
                    with col_env1:
                        st.metric("Fuel Needed", f"{fuel_consumption:.1f}L")
                    with col_env2:
                        st.metric("CO2 Emissions", f"{co2_emissions:.1f}kg")
                    with col_env3:
                        tree_offset = co2_emissions / 21.77  # kg CO2 absorbed by one tree per year
                        st.metric("Trees to Offset", f"{tree_offset:.1f}")
                    with col_env4:
                        fuel_cost = fuel_consumption * 100  # Assuming ₹100 per liter
                        st.metric("Fuel Cost", f"₹{fuel_cost:.0f}")
                
                # Detailed report generation
                st.markdown("### Generate Reports")
                
                col_rep1, col_rep2 = st.columns(2)
                
                with col_rep1:
                    if st.button("📄 Generate Markdown Report", use_container_width=True):
                        report_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        weather_condition = st.session_state.weather_data["condition"]
                        
                        fuel_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
                            best_path, network, vehicle_type, fuel_type
                        )
                        
                        report_content = f"""
# Punjab Route Optimization Report

**Generated:** {report_time}  
**Route:** {source_city} → {destination_city}  
**Vehicle:** {vehicle_type}  
**Fuel Type:** {fuel_type}  
**Optimization:** {preference} Priority  
**Weather:** {weather_condition}  

## Route Summary

### Optimal Route Analysis
- **Best Algorithm:** {best_distance_algo}
- **Total Distance:** {best_distance:.1f} km
- **Estimated Travel Time:** {best_result['travel_time']:.0f} minutes
- **Fuel Consumption:** {fuel_consumption:.2f} L
- **Total Cost Estimate:** ₹{best_result['cost']:.0f}
- **CO2 Emissions:** {co2_emissions:.2f} kg
- **Route Efficiency:** {efficiency:.1f}%

### Algorithm Performance
"""
                        
                        for algo_name, data in results.items():
                            if data.get('distance', float('inf')) != float('inf'):
                                report_content += f"""
#### {algo_name}
- **Execution Time:** {data['time']:.3f} ms
- **Route Distance:** {data['distance']:.1f} km
- **Travel Cost:** ₹{data.get('cost', 0):.0f}
- **Travel Time:** {data.get('travel_time', 0):.0f} minutes
- **Fuel Consumption:** {data.get('fuel_consumption', 0):.2f} L
"""
                        
                        report_content += f"""
### Environmental Impact
- **Fuel Type:** {fuel_type}
- **Fuel Consumption:** {fuel_consumption:.2f} liters
- **CO2 Emissions:** {co2_emissions:.2f} kg
- **Trees to Offset:** {co2_emissions / 21.77:.1f}
- **Environmental Rating:** {'Eco-Friendly' if co2_emissions < 50 else 'Moderate Impact'}

---
*Report generated by SmartRoute Pro - Punjab Edition*
"""
                        
                        st.download_button(
                            label="📥 Download Markdown Report",
                            data=report_content,
                            file_name=f"route_report_{source_city}_{destination_city}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )
                
                with col_rep2:
                    if st.button("📝 Generate Word Report", use_container_width=True):
                        with st.spinner("Generating comprehensive Word document..."):
                            word_doc = generate_word_report(
                                source_city, destination_city, results, vehicle_type,
                                fuel_type, preference, st.session_state.weather_data,
                                st.session_state.traffic_data, best_path, network, pathfinder
                            )
                            
                            st.download_button(
                                label="📥 Download Word Report",
                                data=word_doc,
                                file_name=f"route_report_{source_city}_{destination_city}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
    
    # Footer with information
    st.markdown("---")
    st.markdown("## Punjab: The Land of Five Rivers")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        ### Geography & Culture
        - **Area:** 50,362 km²
        - **Population:** 27.7 million
        - **Capital:** Chandigarh
        - **Languages:** Punjabi, Hindi, English
        """)
    
    with col2:
        st.markdown("""
        ### Transportation Network
        - **National Highways:** 1,500+ km
        - **State Highways:** 1,800+ km
        - **Rural Roads:** 60,000+ km
        - **Railway Network:** 2,200+ km
        """)
    
    with col3:
        st.markdown("""
        ### Economy & Industry
        - **Agriculture:** Wheat, Rice, Cotton
        - **Industries:** Textiles, Sports goods
        - **GDP:** ₹5.6 lakh crore (2023)
        - **Per Capita Income:** ₹1.8 lakh
        """)
    
    with col4:
        st.markdown("""
        ### Network Stats
        - **Cities Covered:** 25 locations
        - **Network Density:** High connectivity
        - **Average Distance:** 65 km between cities
        - **Algorithms:** 4 optimization methods
        """)
def generate_word_report(source: str, destination: str, results: Dict, vehicle_type: str, 
                        fuel_type: str, preference: str, weather_data: Dict, 
                        traffic_data: Dict, best_path: List[str], network: Dict,
                        pathfinder: AdvancedPunjabPathfinder) -> bytes:
    """Generate comprehensive Word document report"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('Punjab Route Optimization Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(255, 107, 53)
    
    # Subtitle
    subtitle = doc.add_heading('SmartRoute Pro - Advanced Multi-Criteria Analysis', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Report metadata
    report_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    metadata_table = doc.add_table(rows=6, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    metadata_data = [
        ('Report Generated', report_time),
        ('Route', f'{source} → {destination}'),
        ('Vehicle Type', f'{VEHICLE_TYPES[vehicle_type]["icon"]} {vehicle_type}'),
        ('Fuel Type', f'{FUEL_TYPES[fuel_type]["icon"]} {fuel_type}'),
        ('Optimization Priority', f'{ROUTE_PREFERENCES[preference]["icon"]} {preference}'),
        ('Weather Condition', f'{weather_data["condition"]} ({weather_data["temperature"]}°C)')
    ]
    
    for i, (label, value) in enumerate(metadata_data):
        row = metadata_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    # Find best algorithm
    valid_results = {k: v for k, v in results.items() if v.get('distance', float('inf')) != float('inf')}
    if valid_results:
        best_algo = min(valid_results.keys(), key=lambda x: valid_results[x]['distance'])
        best_result = valid_results[best_algo]
        
        summary_text = (
            f"The optimal route from {source} to {destination} was calculated using {len(results)} "
            f"pathfinding algorithms. The {best_algo} algorithm provided the most efficient route "
            f"with a total distance of {best_result['distance']:.1f} km, estimated travel time of "
            f"{best_result['travel_time']:.0f} minutes, and total cost of ₹{best_result['cost']:.0f}."
        )
        doc.add_paragraph(summary_text)
        
        # Key metrics table
        doc.add_heading('Key Performance Metrics', level=2)
        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light List Accent 1'
        
        metrics_data = [
            ('Best Algorithm', best_algo),
            ('Total Distance', f"{best_result['distance']:.1f} km"),
            ('Estimated Travel Time', f"{best_result['travel_time']:.0f} minutes"),
            ('Total Fuel Cost', f"₹{best_result['cost']:.0f}"),
            ('Fuel Consumption', f"{best_result.get('fuel_consumption', 0):.2f} L")
        ]
        
        for i, (metric, value) in enumerate(metrics_data):
            row = metrics_table.rows[i]
            row.cells[0].text = metric
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Route Details
        doc.add_heading('Detailed Route Information', level=1)
        
        if best_path:
            doc.add_heading('Step-by-Step Route', level=2)
            route_table = doc.add_table(rows=len(best_path) + 1, cols=6)
            route_table.style = 'Light Grid Accent 1'
            
            # Header row
            headers = ['Step', 'City', 'Type', 'Population', 'Distance', 'Cumulative']
            for i, header in enumerate(headers):
                cell = route_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Route data
            cumulative_distance = 0
            for idx, city in enumerate(best_path):
                city_info = PUNJAB_CITIES[city]
                row = route_table.rows[idx + 1]
                
                if idx == 0:
                    row.cells[0].text = str(idx + 1)
                    row.cells[1].text = city
                    row.cells[2].text = city_info["type"].title()
                    row.cells[3].text = f"{city_info['population']:,}"
                    row.cells[4].text = "START"
                    row.cells[5].text = "0 km"
                else:
                    segment_dist = network.get(best_path[idx-1], {}).get(city, 0)
                    cumulative_distance += segment_dist
                    
                    row.cells[0].text = str(idx + 1)
                    row.cells[1].text = city
                    row.cells[2].text = city_info["type"].title()
                    row.cells[3].text = f"{city_info['population']:,}"
                    row.cells[4].text = f"{segment_dist:.1f} km"
                    row.cells[5].text = f"{cumulative_distance:.1f} km"
        
        doc.add_paragraph()
        
        # Algorithm Performance Comparison
        doc.add_heading('Algorithm Performance Analysis', level=1)
        
        algo_table = doc.add_table(rows=len(results) + 1, cols=5)
        algo_table.style = 'Medium Shading 1 Accent 1'
        
        # Headers
        algo_headers = ['Algorithm', 'Execution Time', 'Distance', 'Cost', 'Travel Time']
        for i, header in enumerate(algo_headers):
            cell = algo_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Algorithm data
        for idx, (algo_name, data) in enumerate(results.items(), 1):
            row = algo_table.rows[idx]
            row.cells[0].text = algo_name
            row.cells[1].text = f"{data['time']:.3f} ms"
            
            if data.get('distance', float('inf')) == float('inf'):
                row.cells[2].text = "No path"
                row.cells[3].text = "-"
                row.cells[4].text = "-"
            else:
                row.cells[2].text = f"{data['distance']:.1f} km"
                row.cells[3].text = f"₹{data.get('cost', 0):.0f}"
                row.cells[4].text = f"{data.get('travel_time', 0):.0f} min"
        
        doc.add_paragraph()
        
        # Fuel and Cost Analysis
        doc.add_heading('Fuel & Cost Analysis', level=1)
        
        fuel_specs = FUEL_TYPES[fuel_type]
        vehicle_specs = VEHICLE_TYPES[vehicle_type]
        
        fuel_cost, fuel_consumption, co2_emissions = pathfinder.calculate_route_cost(
            best_path, network, vehicle_type, fuel_type
        )
        
        fuel_table = doc.add_table(rows=7, cols=2)
        fuel_table.style = 'Light List Accent 1'
        
        fuel_data = [
            ('Fuel Type', fuel_type),
            ('Price per Liter/kg', f"₹{fuel_specs['price_per_liter']:.2f}"),
            ('Vehicle Efficiency', f"{vehicle_specs['fuel_efficiency']} km/L"),
            ('Fuel Consumption', f"{fuel_consumption:.2f} {'L' if fuel_type != 'Electric' else 'kWh'}"),
            ('Total Fuel Cost', f"₹{fuel_cost:.2f}"),
            ('CO2 Emissions', f"{co2_emissions:.2f} kg"),
            ('Environmental Rating', 'Eco-Friendly' if co2_emissions < 50 else 'Moderate Impact')
        ]
        
        for i, (label, value) in enumerate(fuel_data):
            row = fuel_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Environmental Impact
        doc.add_heading('Environmental Impact Assessment', level=1)
        
        tree_offset = co2_emissions / 21.77  # kg CO2 absorbed by one tree per year
        
        env_text = (
            f"This journey will produce approximately {co2_emissions:.2f} kg of CO2 emissions. "
            f"To offset these emissions, approximately {tree_offset:.1f} trees would need to be "
            f"planted and grown for one year. "
        )
        
        if fuel_type == "Electric":
            env_text += "Using electric vehicles significantly reduces emissions compared to conventional fuels."
        elif fuel_type == "CNG":
            env_text += "CNG is a cleaner alternative to petrol and diesel, producing fewer emissions."
        
        doc.add_paragraph(env_text)
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        
        recommendations = []
        
        if fuel_type in ["Petrol", "Diesel"]:
            recommendations.append("Consider switching to CNG or Electric vehicles for reduced emissions and lower operating costs.")
        
        if best_result['travel_time'] > 180:
            recommendations.append("For long journeys, plan for rest stops every 2-3 hours for safety.")
        
        if weather_data["condition"] in ["Heavy Rain", "Fog"]:
            recommendations.append(f"Current weather ({weather_data['condition']}) may affect travel time. Drive carefully and allow extra time.")
        
        # Check traffic in route cities
        high_traffic_cities = []
        for city in best_path:
            if city in traffic_data and traffic_data[city]["level"] == "High":
                high_traffic_cities.append(city)
        
        if high_traffic_cities:
            recommendations.append(f"High traffic expected in: {', '.join(high_traffic_cities)}. Consider alternative departure times.")
        
        for rec in recommendations:
            p = doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run('Report generated by SmartRoute Pro - Punjab Edition')
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()
if __name__ == "__main__":
    main()